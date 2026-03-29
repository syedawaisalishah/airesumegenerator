from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import textwrap

# load_dotenv()

st.set_page_config(page_title="ATS Resume Generator", layout="wide")
st.title("ATS Resume Generator")

llm = ChatGroq(model="openai/gpt-oss-120b", temperature=0)

prompt = """
You are an expert ATS resume writer, technical recruiter, hiring manager, and career strategist.

Rewrite and optimize my resume based on:
1. my current resume text
2. the target job description

Rules:
• Do not invent fake experience. Keep it truthful.
• Output a HIGHLY PROFESSIONAL, ATS-FRIENDLY, RECRUITER-READY resume.
• Only return the final resume. Do NOT include analysis, tables, or extra explanations.
• Remove all section headings and replace them with actual content titles as plain bold text.
  - Do NOT use Markdown symbols (** or __) for bold. Just output as plain bold titles.
• Use ONLY bullet points (•) for lists. No hyphens (-) or other symbols.
• Keep the formatting simple, consistent, and professional.
• Include sections: Professional Summary, Core Skills, Experience, Projects, Education, Certifications, Additional Information.
• Include dates, locations, and details as in the original text.

Output format:
FULL NAME (bold)
Phone | Email | LinkedIn | GitHub | Portfolio | Location

PROFESSIONAL TITLE (bold)

Professional Summary (bold)
• Bullet points here

Core Skills (bold)
• Bullet points here

Experience (bold)
• Bullet points here

Projects (bold)
• Bullet points here

Education (bold)
• Bullet points here

Certifications (bold)
• Bullet points here

Additional Information (bold)
• Bullet points here

======================
MY CURRENT RESUME TEXT
======================
{resume_content}

======================
TARGET JOB DESCRIPTION
======================
{job_description_content}
"""

resume_content = st.text_area("Enter your resume content here", height=300)
job_description_content = st.text_area("Enter your job description content here", height=300)

template = ChatPromptTemplate.from_messages([
    ("system", prompt)
])

# -------- DOCX FORMATTER FUNCTION --------
def create_resume_docx(resume_text):
    with BytesIO() as buffer:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)

        lines = resume_text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            line_clean = line.replace("**", "").replace("__", "")

            # Full name (first line usually)
            if len(line_clean.split()) <= 6 and line_clean.isupper():
                p = doc.add_paragraph()
                run = p.add_run(line_clean)
                run.bold = True
                run.font.size = Pt(16)
            # Section titles
            elif line_clean.upper() in [
                "PROFESSIONAL TITLE",
                "PROFESSIONAL SUMMARY",
                "CORE SKILLS",
                "PROFESSIONAL EXPERIENCE",
                "PROJECTS",
                "EDUCATION",
                "CERTIFICATIONS",
                "ADDITIONAL INFORMATION"
            ] or line_clean.istitle():
                p = doc.add_paragraph()
                run = p.add_run(line_clean)
                run.bold = True
                run.font.size = Pt(13)
            # Bullet points
            elif line_clean.startswith("• "):
                doc.add_paragraph(line_clean[2:], style="List Bullet")
            else:
                p = doc.add_paragraph()
                run = p.add_run(line_clean)
                run.font.size = Pt(11)

        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()  # return bytes for Streamlit download

# -------- PDF FORMATTER FUNCTION WITH WRAPPING --------
def create_resume_pdf(resume_text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 50
    y = height - 50
    max_width = width - 2 * margin
    line_height = 15

    lines = resume_text.split("\n")
    wrapper = textwrap.TextWrapper(width=95)

    for line in lines:
        line_clean = line.replace("**", "").replace("__", "").strip()
        if not line_clean:
            y -= 5
            continue

        # Bold title lines
        if line_clean.isupper() or line_clean.istitle():
            c.setFont("Helvetica-Bold", 12)
            wrapped = wrapper.wrap(text=line_clean)
            for wline in wrapped:
                c.drawString(margin, y, wline)
                y -= line_height
        # Bullet points
        elif line_clean.startswith("• "):
            c.setFont("Helvetica", 11)
            wrapped = wrapper.wrap(text="• " + line_clean[2:])
            for wline in wrapped:
                c.drawString(margin + 10, y, wline)
                y -= line_height
        # Normal text
        else:
            c.setFont("Helvetica", 11)
            wrapped = wrapper.wrap(text=line_clean)
            for wline in wrapped:
                c.drawString(margin, y, wline)
                y -= line_height

        # New page if needed
        if y < 50:
            c.showPage()
            y = height - 50

    c.save()
    buffer.seek(0)
    return buffer

# -------- GENERATE BUTTON --------
if st.button("Generate Resume"):
    if not resume_content.strip() or not job_description_content.strip():
        st.warning("Please enter both resume and job description.")
    else:
        with st.spinner("Generating resume..."):
            formatted_prompt = template.format_messages(
                resume_content=resume_content,
                job_description_content=job_description_content
            )

            response = llm.invoke(formatted_prompt)
            st.session_state.generated_resume = response.content

# -------- SHOW RESULT --------
if "generated_resume" in st.session_state:
    st.subheader("Generated Resume")
    st.text_area("Resume Output", st.session_state.generated_resume, height=500)

    # DOCX download
    docx_bytes = create_resume_docx(st.session_state.generated_resume)
    st.download_button(
        label="Download Resume as DOCX",
        data=docx_bytes,
        file_name="ATS_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # TXT download
    st.download_button(
        label="Download Resume as TXT",
        data=st.session_state.generated_resume,
        file_name="ATS_Resume.txt",
        mime="text/plain"
    )

    # PDF download
    pdf_file = create_resume_pdf(st.session_state.generated_resume)
    st.download_button(
        label="Download Resume as PDF",
        data=pdf_file,
        file_name="ATS_Resume.pdf",
        mime="application/pdf"
    )
